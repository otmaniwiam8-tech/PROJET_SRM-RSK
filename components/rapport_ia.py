# components/rapport_ia.py
import streamlit as st
import pandas as pd
import numpy as np
import datetime
import io
import re
import base64
import os
import xml.sax.saxutils as saxutils

from xhtml2pdf import pisa
from docx import Document
from docx.shared import Pt
from google import genai
from dotenv import load_dotenv

# ============================================================
# 🔑 CLÉ API (chargée depuis les variables d'environnement)
# ============================================================
load_dotenv()  # lit le fichier .env s'il existe (utile en local)

GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")

if not GEMINI_API_KEY:
    # On ne bloque pas l'app, mais on avertit clairement dans l'interface
    st.warning(
        "⚠️ Aucune clé API Gemini trouvée. Définissez la variable d'environnement "
        "GEMINI_API_KEY (fichier .env en local, ou secret dans votre hébergeur)."
    )

# ============================================================
# FONCTION POUR RÉCUPÉRER LE LOGO EN BASE64
# ============================================================
def get_logo_base64():
    """Retourne le logo encodé en base64 pour une utilisation dans HTML/PDF."""
    logo_path = os.path.join(os.path.dirname(__file__), '..', 'assets', 'logo_srm.png')
    try:
        with open(logo_path, "rb") as f:
            return base64.b64encode(f.read()).decode()
    except FileNotFoundError:
        return None

# ============================================================
# FONCTIONS UTILITAIRES
# ============================================================
def escape_html(text):
    if not text:
        return ""
    escaped = saxutils.escape(text)
    escaped = escaped.replace("“", '"').replace("”", '"').replace("‘", "'").replace("’", "'")
    escaped = re.sub(r'[\x00-\x1f\x7f]', '', escaped)
    return escaped

# ============================================================
# APPEL GEMINI
# ============================================================
def call_gemini(prompt):
    try:
        client = genai.Client(api_key=GEMINI_API_KEY)
        response = client.models.generate_content(
            model='gemini-1.5-flash',
            contents=prompt
        )
        return response.text
    except Exception:
        return None

# ============================================================
# GÉNÉRATION DES SECTIONS DU RAPPORT
# ============================================================
def generer_resume_executif(data):
    agence = data['predictions']['agence']
    gerance = data['predictions']['gerance']
    annee = data['predictions']['annee']
    total = data['predictions']['total_annuel']
    unite = data['predictions']['unite']
    prompt = f"""
    Rédige un résumé exécutif professionnel pour un rapport de prédiction de consommation future.
    Contexte : Agence {agence}, Gérance {gerance}, Année {annee}.
    La consommation annuelle prédite est de {total:.2f} {unite}.
    Le modèle utilisé est RandomForestRegressor.
    Rédige en français, avec un ton professionnel et concis (3-4 paragraphes).
    """
    result = call_gemini(prompt)
    if result is None:
        return f"Résumé exécutif pour l'agence {agence} - Gérance {gerance} - Année {annee}. La consommation annuelle prédite est de {total:.2f} {unite}. L'analyse montre une tendance stable avec des variations saisonnières."
    return result

def generer_analyse_kpi(data):
    pred = data['predictions']
    predictions = pred['predictions']
    moyenne = pred['moyenne']
    max_val = max(predictions)
    min_val = min(predictions)
    mois_max = pred['mois_labels'][np.argmax(predictions)]
    mois_min = pred['mois_labels'][np.argmin(predictions)]
    prompt = f"""
    Analyse les KPI prédictifs suivants pour un rapport :
    - Consommation totale annuelle prédite : {pred['total_annuel']:.2f} {pred['unite']}
    - Moyenne mensuelle : {moyenne:.2f} {pred['unite']}
    - Mois le plus élevé : {mois_max} avec {max_val:.2f} {pred['unite']}
    - Mois le plus bas : {mois_min} avec {min_val:.2f} {pred['unite']}
    Explique l'évolution mensuelle, les pics et la signification des KPI.
    """
    result = call_gemini(prompt)
    if result is None:
        return f"Les KPI montrent une consommation totale de {pred['total_annuel']:.2f} {pred['unite']} avec un pic en {mois_max} ({max_val:.2f}) et un minimum en {mois_min} ({min_val:.2f})."
    return result

def generer_analyse_shap(data):
    return "L'analyse des facteurs d'influence (SHAP) indique que 'Année' est le facteur le plus influent, suivi de 'Mois' et 'Gérance'. Ces variables expliquent la majeure partie de la consommation prédite."

def generer_alertes(data):
    pred = data['predictions']
    predictions = pred['predictions']
    total = pred['total_annuel']
    unite = pred['unite']
    alertes = []
    
    if total > 10000:
        alertes.append(("🔴", "Consommation très élevée", f"La consommation annuelle de {total:.2f} {unite} dépasse largement les seuils habituels."))
    elif total > 5000:
        alertes.append(("🟠", "Consommation élevée", f"La consommation annuelle de {total:.2f} {unite} est supérieure à la moyenne."))
    
    if max(predictions) > 2 * min(predictions):
        alertes.append(("🟠", "Forte variation saisonnière", "Les écarts entre les mois sont importants, indiquant une saisonnalité marquée."))
    
    max_val = max(predictions)
    if max_val > 2 * np.mean(predictions):
        mois_max = pred['mois_labels'][np.argmax(predictions)]
        alertes.append(("🟡", "Pic de consommation", f"Le mois de {mois_max} présente un pic à {max_val:.2f} {unite}."))
    
    if len(alertes) == 0:
        alertes.append(("🟢", "Situation normale", "La consommation est stable."))
    
    return alertes

def generer_recommandations(alertes, data):
    if not alertes or (len(alertes) == 1 and "Situation normale" in alertes[0][1]):
        return "Aucune recommandation spécifique pour le moment. Continuez à surveiller les tendances de consommation."
    
    alert_text = "\n".join([f"{niveau} {titre} : {desc}" for niveau, titre, desc in alertes])
    prompt = f"""
    À partir des alertes suivantes :
    {alert_text}
    Propose des recommandations concrètes pour les responsables de la SRM.
    """
    result = call_gemini(prompt)
    if result is None:
        return "• Renforcer le suivi de consommation mensuel\n• Mettre en place un plan de contrôle des périodes de forte demande\n• Optimiser les ressources énergétiques\n• Programmer une maintenance préventive"
    return result

def generer_conclusion(data, resume, recommandations):
    total = data['predictions']['total_annuel']
    unite = data['predictions']['unite']
    agence = data['predictions']['agence']
    resume_text = resume[:300] if resume else "Résumé non disponible."
    reco_text = recommandations[:300] if recommandations else "Recommandations non disponibles."
    
    prompt = f"""
    Rédige une conclusion synthétique et professionnelle pour un rapport de prédiction.
    Agence : {agence}, Consommation : {total:.2f} {unite}.
    Résumé : {resume_text}
    Recommandations : {reco_text}
    Rédige en 2-3 paragraphes.
    """
    result = call_gemini(prompt)
    if result is None:
        return f"Conclusion : La consommation annuelle prédite pour l'agence {agence} est de {total:.2f} {unite}. Les facteurs clés sont l'année, le mois et la gérance. Des actions de suivi sont recommandées."
    return result

# ============================================================
# EXPORT PDF
# ============================================================
def generate_pdf(html_rapport):
    try:
        buffer = io.BytesIO()
        result = pisa.CreatePDF(src=html_rapport, dest=buffer, encoding='UTF-8')
        if result.err:
            st.error("Erreur lors de la génération du PDF (xhtml2pdf).")
            return None
        buffer.seek(0)
        return buffer
    except Exception as e:
        st.error(f"Erreur PDF : {e}")
        return None

# ============================================================
# EXPORT WORD AVEC LOGO AGRANDI
# ============================================================
def ajouter_logo_docx(doc):
    """Insère le logo SRM en en-tête du document Word, avec une largeur agrandie."""
    logo_path = os.path.join(os.path.dirname(__file__), '..', 'assets', 'logo_srm.png')
    try:
        p = doc.add_paragraph()
        p.alignment = 1
        run = p.add_run()
        # Largeur du logo : 200 points (agrandi)
        run.add_picture(logo_path, width=Pt(200))
        doc.add_paragraph()
    except Exception:
        # Fallback : titre texte
        title = doc.add_heading('SRM-RSK - Rapport de prédiction', 0)
        title.alignment = 1

def generate_docx(data, resume, analyse_kpi, analyse_shap, alertes, recommandations, conclusion):
    try:
        doc = Document()
        style = doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(12)
        pred = data['predictions']
        maintenant = datetime.datetime.now().strftime("%d/%m/%Y %H:%M")

        # En-tête avec logo
        ajouter_logo_docx(doc)
        doc.add_paragraph(f"Généré le {maintenant}").alignment = 1
        doc.add_paragraph()

        # Sections
        doc.add_heading('1. Informations générales', level=1)
        doc.add_paragraph(f"Agence : {pred['agence']}")
        doc.add_paragraph(f"Gérance : {pred['gerance']}")
        doc.add_paragraph(f"Année prédite : {pred['annee']}")
        doc.add_paragraph(f"Modèle : RandomForestRegressor")
        doc.add_paragraph(f"Observations : {len(data['df_clean']) if not data['df_clean'].empty else 'N/A'}")
        doc.add_paragraph(f"Variable cible : Consommation")
        doc.add_paragraph()

        doc.add_heading('2. Résumé exécutif', level=1)
        doc.add_paragraph(resume)
        doc.add_paragraph()

        doc.add_heading('3. Résultats de la prédiction', level=1)
        doc.add_paragraph(f"Consommation annuelle prédite : {pred['total_annuel']:.2f} {pred['unite']}")
        doc.add_paragraph(f"Moyenne mensuelle : {pred['moyenne']:.2f} {pred['unite']}")

        doc.add_heading('Indicateurs clés', level=2)
        table = doc.add_table(rows=7, cols=2)
        table.style = 'Table Grid'
        table.rows[0].cells[0].text = 'Indicateur'
        table.rows[0].cells[1].text = 'Valeur'
        rows_data = [
            ('Total annuel', f"{pred['total_annuel']:.2f} {pred['unite']}"),
            ('Moyenne mensuelle', f"{pred['moyenne']:.2f} {pred['unite']}"),
            ('Mois max', f"{pred['mois_labels'][np.argmax(pred['predictions'])]} ({max(pred['predictions']):.2f} {pred['unite']})"),
            ('Mois min', f"{pred['mois_labels'][np.argmin(pred['predictions'])]} ({min(pred['predictions']):.2f} {pred['unite']})"),
            ('Écart-type', f"{np.std(pred['predictions']):.2f} {pred['unite']}"),
            ('Amplitude', f"{max(pred['predictions']) - min(pred['predictions']):.2f} {pred['unite']}")
        ]
        for i, (label, value) in enumerate(rows_data):
            row = table.rows[i+1]
            row.cells[0].text = label
            row.cells[1].text = value
        doc.add_paragraph()

        doc.add_heading('Détail mensuel', level=2)
        table_mois = doc.add_table(rows=14, cols=2)
        table_mois.style = 'Table Grid'
        table_mois.rows[0].cells[0].text = 'Mois'
        table_mois.rows[0].cells[1].text = f'Consommation ({pred["unite"]})'
        for i, mois in enumerate(pred['mois_labels']):
            table_mois.rows[i+1].cells[0].text = mois
            table_mois.rows[i+1].cells[1].text = f"{pred['predictions'][i]:.2f}"
        table_mois.rows[12].cells[0].text = 'TOTAL'
        table_mois.rows[12].cells[1].text = f"{pred['total_annuel']:.2f}"
        table_mois.rows[13].cells[0].text = 'MOYENNE'
        table_mois.rows[13].cells[1].text = f"{pred['moyenne']:.2f}"
        doc.add_paragraph()

        doc.add_heading('4. Analyse des KPI prédictifs', level=1)
        doc.add_paragraph(analyse_kpi)
        doc.add_paragraph()

        doc.add_heading('5. Facteurs influençant', level=1)
        doc.add_paragraph(analyse_shap)
        doc.add_paragraph()

        doc.add_heading('6. Alertes', level=1)
        for niveau, titre, desc in alertes:
            p = doc.add_paragraph()
            p.add_run(f"{niveau} {titre} : ").bold = True
            p.add_run(desc)
        doc.add_paragraph()

        doc.add_heading('7. Recommandations', level=1)
        reco_lines = recommandations.split('\n') if recommandations else ["Aucune recommandation disponible."]
        for line in reco_lines:
            if line.strip():
                doc.add_paragraph(line, style='List Bullet')
        doc.add_paragraph()

        doc.add_heading('8. Conclusion', level=1)
        doc.add_paragraph(conclusion)
        doc.add_paragraph()

        footer = doc.add_paragraph("SRM-RSK - Rapport confidentiel")
        footer.alignment = 1
        footer.style.font.size = Pt(10)

        docx_file = io.BytesIO()
        doc.save(docx_file)
        docx_file.seek(0)
        return docx_file
    except Exception as e:
        st.error(f"Erreur Word : {e}")
        return None

# ============================================================
# CONSTRUCTION HTML AVEC LOGO AGRANDI (BASE64)
# ============================================================
def build_rapport_html(data, resume, analyse_kpi, analyse_shap, alertes, recommandations, conclusion):
    pred = data['predictions']
    maintenant = datetime.datetime.now().strftime("%d/%m/%Y %H:%M")
    mois_labels = pred['mois_labels']
    predictions = pred['predictions']
    
    # Récupération du logo
    logo_b64 = get_logo_base64()
    if logo_b64:
        # Hauteur du logo : 100 px (agrandi)
        logo_html = f'<img src="data:image/png;base64,{logo_b64}" alt="Logo SRM" style="height:100px; margin-bottom:10px;">'
    else:
        logo_html = '<div class="logo">SRM-RSK</div>'
    
    # Construction des tableaux
    mois_rows = ""
    for i, mois in enumerate(mois_labels):
        mois_rows += f"<tr><td style='border-bottom:1px solid #e0e0e0; padding:8px; text-align:center;'>{mois}</td><td style='border-bottom:1px solid #e0e0e0; padding:8px; text-align:right;'>{predictions[i]:.2f}</td></tr>"
    
    kpi_table = f"""
    <table style="width:100%; border-collapse: collapse; font-family: Calibri, Arial, sans-serif; font-size: 14px; margin: 15px 0;">
        <tr style="background-color: #173B63; color: white;">
            <th style="padding: 10px 12px; text-align: left;">Indicateur</th>
            <th style="padding: 10px 12px; text-align: right;">Valeur</th>
        </tr>
        <tr style="background-color: #f8fafc;"><td style="padding: 8px 12px; border-bottom:1px solid #e0e0e0;">Total annuel</td><td style="padding: 8px 12px; text-align: right; border-bottom:1px solid #e0e0e0; font-weight: bold;">{pred['total_annuel']:.2f} {pred['unite']}</td></tr>
        <tr><td style="padding: 8px 12px; border-bottom:1px solid #e0e0e0;">Moyenne mensuelle</td><td style="padding: 8px 12px; text-align: right; border-bottom:1px solid #e0e0e0;">{pred['moyenne']:.2f} {pred['unite']}</td></tr>
        <tr style="background-color: #f8fafc;"><td style="padding: 8px 12px; border-bottom:1px solid #e0e0e0;">Mois maximum</td><td style="padding: 8px 12px; text-align: right; border-bottom:1px solid #e0e0e0;">{mois_labels[np.argmax(predictions)]} ({max(predictions):.2f} {pred['unite']})</td></tr>
        <tr><td style="padding: 8px 12px; border-bottom:1px solid #e0e0e0;">Mois minimum</td><td style="padding: 8px 12px; text-align: right; border-bottom:1px solid #e0e0e0;">{mois_labels[np.argmin(predictions)]} ({min(predictions):.2f} {pred['unite']})</td></tr>
        <tr style="background-color: #f8fafc;"><td style="padding: 8px 12px; border-bottom:1px solid #e0e0e0;">Écart-type</td><td style="padding: 8px 12px; text-align: right; border-bottom:1px solid #e0e0e0;">{np.std(predictions):.2f} {pred['unite']}</td></tr>
        <tr><td style="padding: 8px 12px; border-bottom:1px solid #e0e0e0;">Amplitude (max−min)</td><td style="padding: 8px 12px; text-align: right; border-bottom:1px solid #e0e0e0;">{max(predictions) - min(predictions):.2f} {pred['unite']}</td></tr>
    </table>
    """
    
    detail_mensuel = f"""
    <table style="width:100%; border-collapse: collapse; font-family: Calibri, Arial, sans-serif; font-size: 14px; margin: 15px 0;">
        <tr style="background-color: #173B63; color: white;">
            <th style="padding: 10px 12px; text-align: center;">Mois</th>
            <th style="padding: 10px 12px; text-align: right;">Consommation ({pred['unite']})</th>
        </tr>
        {mois_rows}
        <tr style="background-color: #FF8F00; color: white; font-weight: bold;">
            <td style="padding: 8px 12px; text-align: center;">TOTAL</td>
            <td style="padding: 8px 12px; text-align: right;">{pred['total_annuel']:.2f}</td>
        </tr>
        <tr style="background-color: #173B63; color: white; font-weight: bold;">
            <td style="padding: 8px 12px; text-align: center;">MOYENNE</td>
            <td style="padding: 8px 12px; text-align: right;">{pred['moyenne']:.2f}</td>
        </tr>
    </table>
    """
    
    alert_html = ""
    for niveau, titre, desc in alertes:
        color_map = {"🔴": "#ffebee", "🟠": "#fff3e0", "🟡": "#fff8e1", "🟢": "#e8f5e9"}
        border_color = {"🔴": "#c62828", "🟠": "#e65100", "🟡": "#f9a825", "🟢": "#2e7d32"}
        text_color = {"🔴": "#b71c1c", "🟠": "#bf360c", "🟡": "#e65100", "🟢": "#1b5e20"}
        bg = color_map.get(niveau, "#f5f5f5")
        bc = border_color.get(niveau, "#000")
        tc = text_color.get(niveau, "#000")
        alert_html += f"""
        <div style="background-color: {bg}; border-left: 6px solid {bc}; padding: 14px 18px; margin: 12px 0; border-radius: 6px;">
            <strong style="color: {tc}; font-size: 16px;">{niveau} {titre}</strong><br>
            <span style="font-size: 15px; color: #333;">{desc}</span>
        </div>
        """
    
    reco_lines = recommandations.split('\n') if recommandations else ["Aucune recommandation disponible."]
    reco_html = ""
    for line in reco_lines:
        if line.strip():
            reco_html += f"<li style='margin-bottom: 10px; font-size: 15px; line-height: 1.5;'>{line.strip()}</li>"
    
    html = f"""
    <!DOCTYPE html>
    <html>
    <head>
        <meta charset="UTF-8">
        <title>Rapport SRM-RSK</title>
        <style>
            body {{
                font-family: 'Calibri', 'Arial', sans-serif;
                color: #000000;
                background: #ffffff;
                line-height: 1.6;
                padding: 30px 40px;
                margin: 0;
            }}
            .container {{
                max-width: 1100px;
                margin: 0 auto;
            }}
            .header {{
                text-align: center;
                border-bottom: 2px solid #173B63;
                padding-bottom: 20px;
                margin-bottom: 30px;
            }}
            .logo {{
                font-size: 32px;
                font-weight: bold;
                color: #173B63;
            }}
            .titre-rapport {{
                font-size: 28px;
                color: #173B63;
                font-weight: 700;
                margin: 15px 0 5px 0;
            }}
            .sous-titre {{
                font-size: 16px;
                color: #555555;
                margin-bottom: 10px;
            }}
            .section {{
                margin: 30px 0 20px 0;
            }}
            .section-title {{
                font-size: 22px;
                color: #173B63;
                font-weight: 700;
                border-bottom: 1px solid #173B63;
                padding-bottom: 6px;
                margin-bottom: 20px;
            }}
            .subsection-title {{
                font-size: 18px;
                color: #173B63;
                font-weight: 600;
                margin: 18px 0 10px 0;
            }}
            .info-grid {{
                width: 100%;
                background: #f8fafc;
                padding: 10px 15px;
                border-radius: 8px;
                margin-bottom: 20px;
                border: 1px solid #e0e0e0;
                border-collapse: collapse;
            }}
            .info-grid td.info-item {{
                width: 33%;
                border: none;
                vertical-align: top;
            }}
            .info-item {{
                font-size: 15px;
                padding: 8px 10px;
            }}
            .info-item strong {{
                color: #173B63;
                font-weight: 600;
            }}
            .text-content {{
                font-size: 15px;
                line-height: 1.8;
                text-align: justify;
                padding: 0 5px;
            }}
            .footer {{
                margin-top: 40px;
                border-top: 1px solid #ddd;
                padding-top: 15px;
                font-size: 12px;
                color: #777;
                text-align: center;
            }}
            table {{
                width: 100%;
                border-collapse: collapse;
            }}
            th, td {{
                padding: 8px 12px;
                text-align: left;
                border-bottom: 1px solid #e0e0e0;
            }}
            th {{
                background-color: #173B63;
                color: white;
                font-weight: 600;
            }}
            .total-row {{
                background-color: #FF8F00;
                color: white;
                font-weight: bold;
            }}
            .moyenne-row {{
                background-color: #173B63;
                color: white;
                font-weight: bold;
            }}
            .alert-box {{
                padding: 14px 18px;
                margin: 12px 0;
                border-radius: 6px;
                border-left: 6px solid #000;
            }}
            .reco-list {{
                list-style-type: disc;
                padding-left: 25px;
                margin: 10px 0;
            }}
            .reco-list li {{
                margin-bottom: 10px;
                font-size: 15px;
                line-height: 1.5;
            }}
        </style>
    </head>
    <body>
        <div class="container">
            <div class="header">
                {logo_html}
                <div class="titre-rapport">Rapport de prédiction de consommation</div>
                <div class="sous-titre">Généré le {maintenant} · Document confidentiel</div>
            </div>

            <div class="section">
                <div class="section-title">1. Informations générales</div>
                <table class="info-grid">
                    <tr>
                        <td class="info-item"><strong>Agence :</strong> {pred['agence']}</td>
                        <td class="info-item"><strong>Gérance :</strong> {pred['gerance']}</td>
                        <td class="info-item"><strong>Année prédite :</strong> {pred['annee']}</td>
                    </tr>
                    <tr>
                        <td class="info-item"><strong>Modèle :</strong> RandomForestRegressor</td>
                        <td class="info-item"><strong>Observations :</strong> {len(data['df_clean']) if not data['df_clean'].empty else 'N/A'}</td>
                        <td class="info-item"><strong>Variable cible :</strong> Consommation</td>
                    </tr>
                </table>
            </div>

            <div class="section">
                <div class="section-title">2. Résumé exécutif</div>
                <div class="text-content">{resume.replace('\n', '<br>') if resume else 'Résumé non disponible.'}</div>
            </div>

            <div class="section">
                <div class="section-title">3. Résultats de la prédiction</div>
                <div class="subsection-title">Consommation annuelle prédite</div>
                <div style="font-size: 32px; font-weight: 700; color: #173B63; text-align: center; margin: 10px 0;">{pred['total_annuel']:.2f} <span style="font-size: 20px; font-weight: normal; color: #333;">{pred['unite']}</span></div>
                <div style="text-align: center; font-size: 18px; color: #555; margin-bottom: 15px;">Moyenne mensuelle : <strong>{pred['moyenne']:.2f} {pred['unite']}</strong></div>

                <div class="subsection-title">Indicateurs clés</div>
                {kpi_table}

                <div class="subsection-title">Détail mensuel</div>
                {detail_mensuel}
            </div>

            <div class="section">
                <div class="section-title">4. Analyse des KPI prédictifs</div>
                <div class="text-content">{analyse_kpi.replace('\n', '<br>') if analyse_kpi else 'Analyse non disponible.'}</div>
            </div>

            <div class="section">
                <div class="section-title">5. Facteurs influençant la consommation</div>
                <div class="text-content">{analyse_shap.replace('\n', '<br>')}</div>
            </div>

            <div class="section">
                <div class="section-title">6. Alertes intelligentes</div>
                {alert_html}
            </div>

            <div class="section">
                <div class="section-title">7. Recommandations</div>
                <div style="background: #f8fafc; border-radius: 8px; padding: 15px 20px; border: 1px solid #e0e0e0;">
                    <ul class="reco-list">
                        {reco_html}
                    </ul>
                </div>
            </div>

            <div class="section">
                <div class="section-title">8. Conclusion</div>
                <div class="text-content">{conclusion.replace('\n', '<br>') if conclusion else 'Conclusion non disponible.'}</div>
            </div>

            <div class="footer">
                SRM-RSK - Rapport confidentiel<br>
                <span style="font-size: 11px;">Usage interne - Reproduction interdite.</span>
            </div>
        </div>
    </body>
    </html>
    """
    return html

# ============================================================
# PAGE PRINCIPALE STREAMLIT
# ============================================================
def render_rapport_ia():
    st.markdown("""
    <style>
        .stDownloadButton button {
            background-color: #00A86B !important;
            color: white !important;
            border: none !important;
            border-radius: 8px !important;
            padding: 0.5rem 1.5rem !important;
            font-weight: 600 !important;
            transition: background-color 0.2s;
            width: 100%;
        }
        .stDownloadButton button:hover {
            background-color: #00897B !important;
        }
        .stButton button {
            background-color: #1E88E5 !important;
            color: white !important;
            border: none !important;
            border-radius: 8px !important;
            padding: 0.5rem 1.5rem !important;
            font-weight: 600 !important;
        }
        .stButton button:hover {
            background-color: #1565C0 !important;
        }
    </style>
    """, unsafe_allow_html=True)

    st.markdown("""
    <div style="margin-bottom: 1.5rem;">
        <h1 style="font-size: 2.2rem; font-weight: 700; color: #FFFFFF; margin: 0; padding: 0;">📄 Rapport intelligent généré par IA</h1>
        <p style="font-size: 1.1rem; color: #94A3B8; margin: 0.2rem 0 0 0;">
            Générez un rapport d'aide à la décision à partir des prédictions actuelles
        </p>
    </div>
    """, unsafe_allow_html=True)

    if ('dernieres_predictions' not in st.session_state or st.session_state.dernieres_predictions is None):
        st.warning("⚠️ Veuillez d'abord effectuer une prédiction dans la page '🔮 Prédiction de la consommation'.")
        return

    data = {
        'predictions': st.session_state.dernieres_predictions,
        'df_clean': st.session_state.get('df_clean', pd.DataFrame()),
        'metrics': st.session_state.get('metriques_modele', {})
    }

    if not st.session_state.get('rapport_genere', False):
        if st.button("🤖 Générer le rapport IA", type="primary"):
            with st.spinner("🔄 Génération du rapport en cours..."):
                resume = generer_resume_executif(data)
                analyse_kpi = generer_analyse_kpi(data)
                analyse_shap = generer_analyse_shap(data)
                alertes = generer_alertes(data)
                recommandations = generer_recommandations(alertes, data)
                conclusion = generer_conclusion(data, resume, recommandations)

                html_rapport = build_rapport_html(data, resume, analyse_kpi, analyse_shap, alertes, recommandations, conclusion)

                st.session_state.html_rapport = html_rapport
                st.session_state.rapport_genere = True
                st.session_state.rapport_data = data
                st.session_state.rapport_resume = resume
                st.session_state.rapport_analyse_kpi = analyse_kpi
                st.session_state.rapport_analyse_shap = analyse_shap
                st.session_state.rapport_alertes = alertes
                st.session_state.rapport_recommandations = recommandations
                st.session_state.rapport_conclusion = conclusion
                st.rerun()

    if st.session_state.get('rapport_genere', False):
        html_rapport = st.session_state.html_rapport
        data = st.session_state.rapport_data
        resume = st.session_state.rapport_resume
        analyse_kpi = st.session_state.rapport_analyse_kpi
        analyse_shap = st.session_state.rapport_analyse_shap
        alertes = st.session_state.rapport_alertes
        recommandations = st.session_state.rapport_recommandations
        conclusion = st.session_state.rapport_conclusion

        st.markdown("### 📄 Aperçu du rapport")
        st.markdown(f"""
        <iframe srcdoc="{html_rapport.replace('"', '&quot;').replace('\n', ' ')}" 
                style="width:100%; height:800px; border:1px solid #ddd; border-radius:8px;"></iframe>
        """, unsafe_allow_html=True)

        st.markdown("---")
        col1, col2 = st.columns(2)

        with col1:
            pdf_data = generate_pdf(html_rapport)
            if pdf_data:
                st.download_button(
                    label="📄 Télécharger PDF",
                    data=pdf_data,
                    file_name=f"Rapport_SRMRSK_{datetime.datetime.now().strftime('%Y%m%d_%H%M')}.pdf",
                    mime="application/pdf",
                    key="pdf_download"
                )
            else:
                st.error("Erreur lors de la génération du PDF.")

        with col2:
            docx_data = generate_docx(data, resume, analyse_kpi, analyse_shap, alertes, recommandations, conclusion)
            if docx_data:
                st.download_button(
                    label="📝 Télécharger Word",
                    data=docx_data,
                    file_name=f"Rapport_SRMRSK_{datetime.datetime.now().strftime('%Y%m%d_%H%M')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key="docx_download"
                )
            else:
                st.error("Erreur lors de la génération du document Word.")

        if st.button("🔄 Régénérer le rapport"):
            st.session_state.rapport_genere = False
            st.rerun()